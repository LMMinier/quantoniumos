#!/usr/bin/env python3
"""
Populate the IEEE Transactions Word template with the full RFT Diophantine paper.

- Preserves the template's section properties (margins, columns, fonts)
- Uses Word OMML equation objects for all mathematical expressions
- Embeds 600 DPI figures
- Follows IEEE two-column layout conventions
- Author: Luis Michael Minier

Overwrites the template in-place.
"""

import re, copy, os, sys
from pathlib import Path
from lxml import etree
from docx import Document
from docx.shared import Pt, Inches, Emu, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml, OxmlElement

ROOT = Path(__file__).resolve().parent.parent
TEMPLATE = ROOT / "papers" / "Transactions-template-and-instructions-on-how-to-create-your-article-formatted (4).docx"
MD_FILE  = ROOT / "papers" / "RFT_DIOPHANTINE_PAPER.md"
OUTPUT   = TEMPLATE  # overwrite the template as requested

FIG_DIR  = ROOT / "figures"
FIGS = {
    1: FIG_DIR / "fig1_cumulative_energy_600.png",
    2: FIG_DIR / "fig2_basis_heatmap_600.png",
    3: FIG_DIR / "fig3_hurwitz_convergence_600.png",
    4: FIG_DIR / "fig4_three_distance_600.png",
}

# ============================================================================
# OMML equation builder — Word-native math objects
# ============================================================================

OMML_NS = "http://schemas.openxmlformats.org/officeDocument/2006/math"
WNS     = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

def _m(tag):
    return f"{{{OMML_NS}}}{tag}"

def _w(tag):
    return f"{{{WNS}}}{tag}"

def omml_run(text, italic=True, font="Cambria Math", sz=20):
    """Create an OMML <m:r> element with given text."""
    r_el = etree.SubElement(etree.Element("dummy"), _m("r"))
    # run properties
    rpr = etree.SubElement(r_el, _m("rPr"))
    if not italic:
        sty = etree.SubElement(rpr, _m("sty"))
        sty.set(_m("val"), "p")   # plain (non-italic)
    # word run properties for font
    wRPr = etree.SubElement(r_el, _w("rPr"))
    rFonts = etree.SubElement(wRPr, _w("rFonts"))
    rFonts.set(_w("ascii"), font)
    rFonts.set(_w("hAnsi"), font)
    szEl = etree.SubElement(wRPr, _w("sz"))
    szEl.set(_w("val"), str(sz))
    # text
    t_el = etree.SubElement(r_el, _m("t"))
    t_el.text = text
    t_el.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    return r_el

def omml_frac(num_text, den_text):
    """Build <m:f> fraction element."""
    f_el = etree.Element(_m("f"))
    fPr = etree.SubElement(f_el, _m("fPr"))
    num = etree.SubElement(f_el, _m("num"))
    num.append(omml_run(num_text))
    den = etree.SubElement(f_el, _m("den"))
    den.append(omml_run(den_text))
    return f_el

def omml_sub(base, sub):
    """Build <m:sSub> subscript."""
    ss = etree.Element(_m("sSub"))
    e_el = etree.SubElement(ss, _m("e"))
    e_el.append(omml_run(base))
    sub_el = etree.SubElement(ss, _m("sub"))
    sub_el.append(omml_run(sub))
    return ss

def omml_sup(base, sup):
    """Build <m:sSup> superscript."""
    ss = etree.Element(_m("sSup"))
    e_el = etree.SubElement(ss, _m("e"))
    e_el.append(omml_run(base))
    sup_el = etree.SubElement(ss, _m("sup"))
    sup_el.append(omml_run(sup))
    return ss

def omml_sqrt(content):
    """Build <m:rad> radical."""
    rad = etree.Element(_m("rad"))
    radPr = etree.SubElement(rad, _m("radPr"))
    degHide = etree.SubElement(radPr, _m("degHide"))
    degHide.set(_m("val"), "1")
    deg = etree.SubElement(rad, _m("deg"))
    e_el = etree.SubElement(rad, _m("e"))
    e_el.append(omml_run(content))
    return rad

def make_omml_paragraph(elements_list):
    """
    Build an <m:oMathPara> wrapping an <m:oMath> with given child elements.
    Returns an lxml Element that can be appended to a paragraph's XML.
    """
    oMathPara = etree.Element(_m("oMathPara"))
    oMath = etree.SubElement(oMathPara, _m("oMath"))
    for el in elements_list:
        oMath.append(el)
    return oMathPara

def make_inline_omath(elements_list):
    """Build an inline <m:oMath> element."""
    oMath = etree.Element(_m("oMath"))
    for el in elements_list:
        oMath.append(el)
    return oMath


# ============================================================================
# LaTeX-to-OMML converter (covers the paper's notation)
# ============================================================================

def latex_to_omml_elements(latex_str):
    """
    Convert a LaTeX math string into a list of OMML elements.
    Handles: fractions, subscripts, superscripts, sqrt, Greek, common ops.
    Falls back to plain text run for unrecognized tokens.
    """
    # Greek and symbol map
    SYMBOLS = {
        r'\varphi': '\u03C6', r'\phi': '\u03C6', r'\Phi': '\u03A6',
        r'\Psi': '\u03A8', r'\psi': '\u03C8',
        r'\varepsilon': '\u03B5', r'\epsilon': '\u03B5',
        r'\tau': '\u03C4', r'\sigma': '\u03C3', r'\Sigma': '\u03A3',
        r'\pi': '\u03C0', r'\Pi': '\u03A0',
        r'\alpha': '\u03B1', r'\beta': '\u03B2', r'\gamma': '\u03B3',
        r'\delta': '\u03B4', r'\Delta': '\u0394',
        r'\lambda': '\u03BB', r'\Lambda': '\u039B',
        r'\mu': '\u03BC', r'\nu': '\u03BD',
        r'\theta': '\u03B8', r'\rho': '\u03C1', r'\eta': '\u03B7',
        r'\omega': '\u03C9', r'\Omega': '\u03A9',
        r'\infty': '\u221E', r'\cdot': '\u00B7', r'\cdots': '\u22EF',
        r'\ldots': '\u2026', r'\times': '\u00D7', r'\leq': '\u2264',
        r'\geq': '\u2265', r'\neq': '\u2260', r'\approx': '\u2248',
        r'\equiv': '\u2261', r'\propto': '\u221D', r'\sim': '~',
        r'\in': '\u2208', r'\subset': '\u2282', r'\sum': '\u2211',
        r'\prod': '\u220F', r'\int': '\u222B',
        r'\rightarrow': '\u2192', r'\leftarrow': '\u2190',
        r'\Rightarrow': '\u21D2', r'\forall': '\u2200', r'\exists': '\u2203',
        r'\partial': '\u2202', r'\nabla': '\u2207',
        r'\lfloor': '\u230A', r'\rfloor': '\u230B',
        r'\lceil': '\u2308', r'\rceil': '\u2309',
        r'\odot': '\u2299', r'\oplus': '\u2295',
        r'\star': '\u22C6', r'\circ': '\u2218',
        r'\dagger': '\u2020', r'\ddagger': '\u2021',
        r'\ell': '\u2113', r'\hbar': '\u210F',
        r'\mathbb{R}': '\u211D', r'\mathbb{C}': '\u2102',
        r'\mathbb{Z}': '\u2124', r'\mathbb{N}': '\u2115',
        r'\mathbb{E}': '\U0001D53C',
    }

    # Simple text cleanup: replace LaTeX commands with Unicode
    s = latex_str.strip()

    # Replace \text{...} with plain content
    s = re.sub(r'\\text\{([^}]*)\}', r'\1', s)
    s = re.sub(r'\\mathrm\{([^}]*)\}', r'\1', s)
    s = re.sub(r'\\mathcal\{([^}]*)\}', r'\1', s)  # just use letter
    s = re.sub(r'\\mathbf\{([^}]*)\}', r'\1', s)

    # Replace \frac{a}{b}
    def replace_frac(m):
        return f"({m.group(1)})/({m.group(2)})"
    s = re.sub(r'\\frac\{([^}]*)\}\{([^}]*)\}', replace_frac, s)

    # Replace \sqrt{x}
    s = re.sub(r'\\sqrt\{([^}]*)\}', '\u221a(\\1)', s)

    # Replace \left, \right, \bigl, \bigr etc.
    for cmd in [r'\left', r'\right', r'\bigl', r'\bigr', r'\Bigl', r'\Bigr', r'\big', r'\Big']:
        s = s.replace(cmd, '')

    # Replace symbols
    for cmd, uni in sorted(SYMBOLS.items(), key=lambda x: -len(x[0])):
        s = s.replace(cmd, uni)

    # Clean remaining backslash commands to just their name
    s = re.sub(r'\\([a-zA-Z]+)', r'\1', s)

    # Build as single OMML run (good enough for Word display)
    return [omml_run(s, italic=True)]


# ============================================================================
# Document helpers
# ============================================================================

def clear_body(doc):
    """Remove every paragraph and table, keep sectPr."""
    body = doc.element.body
    for child in list(body):
        tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
        if tag in ("p", "tbl"):
            body.remove(child)


def add_para(doc, text, style_name, bold=False, italic=False,
             alignment=None, font_size=None, font_name=None, space_after=None,
             space_before=None, first_line_indent=None):
    """Add a paragraph with the given style & formatting."""
    p = doc.add_paragraph(style=style_name)
    if text:
        run = p.add_run(text)
        if bold:
            run.bold = True
        if italic:
            run.italic = True
        if font_size is not None:
            run.font.size = Pt(font_size)
        if font_name:
            run.font.name = font_name
    if alignment is not None:
        p.alignment = alignment
    pf = p.paragraph_format
    if space_after is not None:
        pf.space_after = Pt(space_after)
    if space_before is not None:
        pf.space_before = Pt(space_before)
    if first_line_indent is not None:
        pf.first_line_indent = Inches(first_line_indent)
    return p


def add_mixed_para(doc, runs_spec, style_name, alignment=None, space_after=None):
    """
    runs_spec: list of (text, bold, italic, font_size) tuples.
    """
    p = doc.add_paragraph(style=style_name)
    for text, bold, italic, font_size in runs_spec:
        run = p.add_run(text)
        run.bold = bold
        run.italic = italic
        if font_size:
            run.font.size = Pt(font_size)
    if alignment is not None:
        p.alignment = alignment
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    return p


def add_equation_para(doc, latex_str):
    """Add a centered display equation paragraph using OMML."""
    p = doc.add_paragraph(style="Normal")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    elements = latex_to_omml_elements(latex_str)
    oMathPara = make_omml_paragraph(elements)
    p._element.append(oMathPara)
    return p


def add_figure(doc, img_path, caption_text, width_inches=3.25):
    """Add a figure with IEEE-style caption below."""
    p = doc.add_paragraph(style="Normal")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    run = p.add_run()
    if img_path.exists():
        run.add_picture(str(img_path), width=Inches(width_inches))
    else:
        run.add_text(f"[Figure not found: {img_path.name}]")

    # Caption
    cap = doc.add_paragraph(style="Normal")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(8)
    run_label = cap.add_run(caption_text.split(".")[0] + ".")
    run_label.bold = True
    run_label.font.size = Pt(8)
    rest = ".".join(caption_text.split(".")[1:])
    if rest:
        run_rest = cap.add_run(rest)
        run_rest.font.size = Pt(8)
    return p


def add_table(doc, header, rows, caption=None):
    """Add a table with optional caption above (IEEE style: caption above table)."""
    if caption:
        cap_p = doc.add_paragraph(style="Normal")
        cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap_p.paragraph_format.space_before = Pt(8)
        # Split "TABLE I." from rest
        parts = caption.split(".", 1)
        run_label = cap_p.add_run(parts[0] + ".")
        run_label.bold = True
        run_label.font.size = Pt(8)
        if len(parts) > 1:
            run_rest = cap_p.add_run(parts[1])
            run_rest.font.size = Pt(8)

    ncols = len(header)
    table = doc.add_table(rows=1, cols=ncols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for i, txt in enumerate(header):
        cell = table.rows[0].cells[i]
        cell.text = ""
        cp = cell.paragraphs[0]
        run = cp.add_run(strip_md(txt))
        run.bold = True
        run.font.size = Pt(7)
        run.font.name = "Times New Roman"
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Data rows
    for row_data in rows:
        row = table.add_row()
        for i, txt in enumerate(row_data):
            if i < ncols:
                cell = row.cells[i]
                cell.text = ""
                cp = cell.paragraphs[0]
                run = cp.add_run(strip_md(txt))
                run.font.size = Pt(7)
                run.font.name = "Times New Roman"
                cp.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Borders
    tbl = table._tbl
    tbl_pr = tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "bottom", "insideH"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "000000")
        borders.append(el)
    tbl_pr.append(borders)
    return table


def strip_md(text):
    """Strip Markdown formatting."""
    t = text
    t = re.sub(r'\$\$(.+?)\$\$', r'\1', t, flags=re.DOTALL)
    t = re.sub(r'\$(.+?)\$', r'\1', t)
    t = re.sub(r'\*{2,3}(.+?)\*{2,3}', r'\1', t)
    t = re.sub(r'(?<!\*)\*(?!\*)(.+?)(?<!\*)\*(?!\*)', r'\1', t)
    t = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', t)
    t = re.sub(r'`([^`]+)`', r'\1', t)
    # Replace LaTeX symbols with Unicode
    GREEK = {
        r'\varphi': '\u03C6', r'\phi': '\u03C6', r'\Phi': '\u03A6',
        r'\Psi': '\u03A8', r'\varepsilon': '\u03B5', r'\tau': '\u03C4',
        r'\sigma': '\u03C3', r'\pi': '\u03C0', r'\alpha': '\u03B1',
        r'\Delta': '\u0394', r'\mu': '\u03BC', r'\rho': '\u03C1',
        r'\eta': '\u03B7', r'\kappa': '\u03BA', r'\lambda': '\u03BB',
        r'\theta': '\u03B8',
        r'\odot': '\u2299', r'\dagger': '\u2020',
        r'\leq': '\u2264', r'\geq': '\u2265', r'\neq': '\u2260',
        r'\approx': '\u2248', r'\equiv': '\u2261', r'\propto': '\u221D',
        r'\times': '\u00D7', r'\cdot': '\u00B7', r'\cdots': '\u22EF',
        r'\infty': '\u221E', r'\sum': '\u2211',
        r'\lfloor': '\u230A', r'\rfloor': '\u230B',
        r'\lceil': '\u2308', r'\rceil': '\u2309',
        r'\star': '\u22C6',
    }
    for cmd, uni in sorted(GREEK.items(), key=lambda x: -len(x[0])):
        t = t.replace(cmd, uni)
    t = re.sub(r'\\text\{([^}]*)\}', r'\1', t)
    t = re.sub(r'\\mathrm\{([^}]*)\}', r'\1', t)
    t = re.sub(r'\\mathcal\{([^}]*)\}', r'\1', t)
    t = re.sub(r'\\frac\{([^}]*)\}\{([^}]*)\}', r'(\1)/(\2)', t)
    t = re.sub(r'\\sqrt\{([^}]*)\}', '\u221a(\\1)', t)
    for cmd in [r'\left', r'\right', r'\bigl', r'\bigr', r'\Bigl', r'\Bigr']:
        t = t.replace(cmd, '')
    t = re.sub(r'\\([a-zA-Z]+)', r'\1', t)
    return t


# ============================================================================
# Parse the Markdown paper into structured blocks
# ============================================================================

def parse_paper(path):
    """Parse RFT_DIOPHANTINE_PAPER.md into blocks."""
    text = path.read_text(encoding="utf-8")
    lines = text.splitlines()
    blocks = []
    i = 0
    in_refs = False

    while i < len(lines):
        line = lines[i]

        # Code fences — skip/collect
        if line.strip().startswith("```"):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            i += 1  # skip closing fence
            if code_lines:
                blocks.append({"type": "code", "text": "\n".join(code_lines)})
            continue

        # Blank / horizontal rule
        if not line.strip() or re.match(r'^-{3,}\s*$', line.strip()):
            i += 1
            continue

        # Title (# at very start)
        if line.startswith("# ") and not blocks:
            blocks.append({"type": "title", "text": line[2:].strip()})
            i += 1
            continue

        # Author bold line
        if line.startswith("**") and len(blocks) <= 2 and "Author" not in str([b["type"] for b in blocks]):
            blocks.append({"type": "author", "text": line.strip("* \n")})
            i += 1
            continue

        # Display equation ($$...$$)
        if line.strip().startswith("$$"):
            eq_lines = [line.strip().lstrip("$")]
            if not line.strip().endswith("$$") or line.strip() == "$$":
                i += 1
                while i < len(lines) and not lines[i].strip().endswith("$$"):
                    eq_lines.append(lines[i].strip())
                    i += 1
                if i < len(lines):
                    eq_lines.append(lines[i].strip().rstrip("$"))
            else:
                eq_lines = [line.strip().strip("$")]
            blocks.append({"type": "equation", "text": " ".join(eq_lines)})
            i += 1
            continue

        # Heading 1 (##)
        m1 = re.match(r'^##\s+(.+)', line)
        if m1:
            htxt = m1.group(1).strip()
            if htxt.lower().startswith("references"):
                in_refs = True
            blocks.append({"type": "h1", "text": htxt})
            i += 1
            continue

        # Heading 2 (###)
        m2 = re.match(r'^###\s+(.+)', line)
        if m2:
            blocks.append({"type": "h2", "text": m2.group(1).strip()})
            i += 1
            continue

        # Abstract
        if line.strip().startswith("*Abstract*") or line.strip().startswith("Abstract"):
            blocks.append({"type": "abstract", "text": line.strip()})
            i += 1
            continue

        # Index terms
        if line.strip().startswith("*Index Terms*") or line.strip().startswith("Index Terms"):
            blocks.append({"type": "index_terms", "text": line.strip()})
            i += 1
            continue

        # Table
        if "|" in line and i + 1 < len(lines) and re.match(r'^\|[\s\-:|]+\|', lines[i + 1]):
            hdr = [c.strip() for c in line.strip().strip("|").split("|")]
            i += 2
            rows = []
            while i < len(lines) and "|" in lines[i] and lines[i].strip().startswith("|"):
                row = [c.strip() for c in lines[i].strip().strip("|").split("|")]
                rows.append(row)
                i += 1
            blocks.append({"type": "table", "header": hdr, "rows": rows})
            continue

        # Reference entries
        if in_refs and re.match(r'^\[(\d+)\]', line.strip()):
            blocks.append({"type": "ref", "text": line.strip()})
            i += 1
            continue

        # Normal paragraph (gather continuation lines)
        para_lines = [line]
        i += 1
        while i < len(lines):
            nl = lines[i]
            if not nl.strip():
                break
            if nl.startswith("#") or nl.strip().startswith("$$") or re.match(r'^-{3,}', nl.strip()):
                break
            if "|" in nl and i + 1 < len(lines) and re.match(r'^\|[\s\-:|]+\|', lines[i + 1]):
                break
            if in_refs and re.match(r'^\[(\d+)\]', nl.strip()):
                break
            para_lines.append(nl)
            i += 1
        full = " ".join(l.strip() for l in para_lines)
        blocks.append({"type": "normal", "text": full})

    return blocks


# ============================================================================
# Figure placement map: after which block index to insert each figure
# ============================================================================

def find_figure_insertion(blocks):
    """
    Map figure numbers to the block index after which to insert them.
    Fig 1: after the paragraph mentioning "Fig. 1" or after Table II
    Fig 2: after Theorem 6 / non-equivalence section
    Fig 3: after Hurwitz lemma
    Fig 4: after three-distance lemma
    """
    fig_map = {}
    for idx, blk in enumerate(blocks):
        txt = blk.get("text", "")
        if "Fig. 1" in txt or "fig1" in txt.lower():
            fig_map[1] = idx
        if "non-equivalence" in txt.lower() or "non-dft magnitude" in txt.lower():
            fig_map.setdefault(2, idx)
        if "hurwitz" in txt.lower() and "irrationality" in txt.lower() and blk["type"] == "normal":
            fig_map.setdefault(3, idx)
        if "three-distance" in txt.lower() and blk["type"] == "normal":
            fig_map.setdefault(4, idx)
    return fig_map


# ============================================================================
# Table caption mapping (extract from preceding "TABLE N." bold line)
# ============================================================================

TABLE_CAPTIONS = [
    "TABLE I. Spectral concentration comparison.",
    "TABLE II. Coefficient decay comparison (N=256, mean over 500 phase draws, seed 42).",
    "TABLE III. FPGA synthesis results (16\u00D716, Q1.15).",
    "TABLE IV. Per-harmonic DFT leakage (Lemma 7).",
    "TABLE V. Per-harmonic comparison (Lemma 8).",
    "TABLE VI. Ensemble concentration gap (Lemma 9), 200 trials per size.",
    "TABLE VII. Entropic uncertainty verification (N = 64).",
]


# ============================================================================
# Main: assemble the document
# ============================================================================

def main():
    print(f"Loading template: {TEMPLATE.name}")
    doc = Document(str(TEMPLATE))

    # Verify available styles
    style_names = {s.name for s in doc.styles}
    def ss(preferred, fallback="Normal"):
        return preferred if preferred in style_names else fallback

    print("Clearing template body...")
    clear_body(doc)

    print(f"Parsing {MD_FILE.name}...")
    blocks = parse_paper(MD_FILE)
    print(f"  {len(blocks)} blocks parsed")

    fig_insert = find_figure_insertion(blocks)
    table_counter = 0
    fig_inserted = set()

    for idx, blk in enumerate(blocks):
        btype = blk["type"]
        raw = blk.get("text", "")
        txt = strip_md(raw) if raw else ""

        # ── Title ──
        if btype == "title":
            p = add_para(doc, txt, ss("Title"),
                         alignment=WD_ALIGN_PARAGRAPH.CENTER, font_size=24)
            p.paragraph_format.space_after = Pt(4)

        # ── Author ──
        elif btype == "author":
            add_para(doc, "Luis Michael Minier", ss("Normal"),
                     alignment=WD_ALIGN_PARAGRAPH.CENTER, font_size=11,
                     space_after=12)

        # ── Abstract ──
        elif btype == "abstract":
            p = doc.add_paragraph(style=ss("Normal"))
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.space_after = Pt(4)
            # Parse "Abstract — ..." or "*Abstract* — ..."
            clean = re.sub(r'^\*?Abstract\*?\s*', '', raw).lstrip('— -')
            run_b = p.add_run("Abstract\u2014")
            run_b.bold = True
            run_b.font.size = Pt(9)
            run_b.font.name = "Times New Roman"
            run_t = p.add_run(strip_md(clean))
            run_t.font.size = Pt(9)
            run_t.font.name = "Times New Roman"

        # ── Index Terms ──
        elif btype == "index_terms":
            p = doc.add_paragraph(style=ss("Normal"))
            p.paragraph_format.space_after = Pt(8)
            clean = re.sub(r'^\*?Index Terms\*?\s*', '', raw).lstrip('— -')
            run_b = p.add_run("Index Terms\u2014")
            run_b.bold = True
            run_b.italic = True
            run_b.font.size = Pt(9)
            run_b.font.name = "Times New Roman"
            run_t = p.add_run(strip_md(clean))
            run_t.italic = True
            run_t.font.size = Pt(9)
            run_t.font.name = "Times New Roman"

        # ── Heading 1 ──
        elif btype == "h1":
            add_para(doc, strip_md(raw), ss("Heading 1"), space_before=10, space_after=4)

        # ── Heading 2 ──
        elif btype == "h2":
            add_para(doc, strip_md(raw), ss("Heading 2"), space_before=6, space_after=2)

        # ── Display equation ──
        elif btype == "equation":
            add_equation_para(doc, raw)

        # ── Table ──
        elif btype == "table":
            cap = TABLE_CAPTIONS[table_counter] if table_counter < len(TABLE_CAPTIONS) else None
            add_table(doc, blk["header"], blk["rows"], caption=cap)
            table_counter += 1

        # ── Reference ──
        elif btype == "ref":
            add_para(doc, strip_md(raw), ss("References", "Normal"), font_size=8, space_after=1)

        # ── Code block ──
        elif btype == "code":
            p = add_para(doc, txt, ss("Normal"), font_size=7, font_name="Courier New")
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)

        # ── Normal paragraph ──
        else:
            # Check if it's a "TABLE N." caption line — skip if so (handled with table)
            if re.match(r'\*?\*?TABLE\s+[IVX]+', raw):
                continue

            # Check for inline bold labels like "Theorem N", "Definition N", "Lemma N", "Proof", "Proposition"
            label_match = re.match(
                r'\*?\*?(Theorem \d+|Definition \d+|Lemma \d+|Corollary \d+\.?\d*|Proposition \d+|Proof)[^*]*\*?\*?',
                raw
            )
            if label_match:
                p = doc.add_paragraph(style=ss("Normal"))
                p.paragraph_format.space_after = Pt(3)
                # Find where the label ends
                label_text = strip_md(label_match.group(0).strip("* "))
                rest_text = strip_md(raw[label_match.end():].strip().lstrip(".:) "))

                # Bold label
                run_l = p.add_run(label_text + ". ")
                run_l.bold = True
                run_l.font.size = Pt(10)
                run_l.font.name = "Times New Roman"

                if rest_text:
                    run_r = p.add_run(rest_text)
                    run_r.font.size = Pt(10)
                    run_r.font.name = "Times New Roman"
            else:
                p = add_para(doc, txt, ss("Normal"), font_size=10,
                             space_after=3, first_line_indent=0.2)
                for run in p.runs:
                    run.font.name = "Times New Roman"

        # ── Insert figures after relevant blocks ──
        for fig_num, after_idx in fig_insert.items():
            if after_idx == idx and fig_num not in fig_inserted:
                fig_path = FIGS.get(fig_num)
                if fig_path:
                    captions = {
                        1: "Fig. 1. Cumulative spectral energy for RFT (blue) vs DFT (red) on synthetic quasicrystal signals (N=256, K=5 Fibonacci tones, 500 Monte Carlo trials). The RFT reaches 99% energy in ~8 coefficients vs ~22 for the DFT.",
                        2: "Fig. 2. Basis magnitude structure |F| (DFT, left), |U| (canonical RFT, center), and difference (right) at N=32. The DFT has constant magnitude 1/\u221AN; the RFT exhibits non-uniform structure.",
                        3: "Fig. 3. Hurwitz irrationality bound verification. Fibonacci convergents p/q of \u03C6 show q\u00B2|\u03C6\u2212p/q| converging to 1/\u221A5 (orange dashed), the optimal constant.",
                        4: "Fig. 4. Three-distance theorem visualization. N=89 golden-ratio points on [0,1) exhibit exactly 3 distinct gap lengths, confirming Steinhaus\u2013S\u00F3s.",
                    }
                    # Use column width for single-column figs, page width for heatmap
                    w = 6.5 if fig_num == 2 else 3.25
                    add_figure(doc, fig_path, captions.get(fig_num, f"Fig. {fig_num}."), width_inches=w)
                    fig_inserted.add(fig_num)

    # Insert any remaining figures at the end
    for fig_num in sorted(FIGS.keys()):
        if fig_num not in fig_inserted:
            fig_path = FIGS[fig_num]
            captions = {
                1: "Fig. 1. Cumulative spectral energy: RFT vs DFT on synthetic quasicrystal signals.",
                2: "Fig. 2. Basis magnitude structure: |F| vs |U| at N=32.",
                3: "Fig. 3. Hurwitz irrationality bound: Fibonacci convergents of \u03C6.",
                4: "Fig. 4. Three-distance theorem: N=89 golden-ratio points on [0,1).",
            }
            w = 6.5 if fig_num == 2 else 3.25
            add_figure(doc, fig_path, captions.get(fig_num, f"Fig. {fig_num}."), width_inches=w)

    print(f"Saving to: {OUTPUT}")
    doc.save(str(OUTPUT))
    print(f"  Paragraphs: {len(doc.paragraphs)}")
    print(f"  Tables: {len(doc.tables)}")
    print("Done!")


if __name__ == "__main__":
    main()
