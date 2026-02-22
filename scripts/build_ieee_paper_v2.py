#!/usr/bin/env python3
"""
Rebuild the IEEE-formatted Diophantine paper DOCX with correct two-column layout.

IEEE template layout:
  - Page: 8.5 x 11 in
  - Margins: 0.65 in each side
  - Two columns, each 3.5 in wide, 0.2 in gutter
  - All images must be <= 3.4 in (fits within column)
  - Tables constrained to column width
  - Body text: 10pt Times New Roman
  - Title spans both columns (already handled by template's Title style)

Author: Luis Michael Minier
"""

import re
from pathlib import Path
from lxml import etree
from docx import Document
from docx.shared import Pt, Inches, Emu, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path(__file__).resolve().parent.parent
TEMPLATE = ROOT / "papers" / "Transactions-template-and-instructions-on-how-to-create-your-article-formatted (4).docx"
MD_FILE  = ROOT / "papers" / "RFT_DIOPHANTINE_PAPER.md"
FIG_DIR  = ROOT / "figures"

# IEEE two-column layout constants
COL_WIDTH_IN = 3.4        # safe image/table width within a 3.5" column
PAGE_WIDTH_IN = 7.0       # full usable width across both columns
BODY_FONT_SIZE = 10       # pt
BODY_FONT = "Times New Roman"
REF_FONT_SIZE = 8
CAPTION_FONT_SIZE = 8
TABLE_FONT_SIZE = 7

FIGS = {
    1: ("fig1_cumulative_energy_600.png",
        "Fig. 1.  Cumulative spectral energy for RFT (blue) vs DFT (red) on "
        "synthetic quasicrystal signals (N = 256, K = 5 Fibonacci tones, "
        "500 Monte Carlo trials). The RFT reaches 99% energy in ~8 coefficients "
        "vs ~22 for the DFT."),
    2: ("fig2_basis_heatmap_600.png",
        "Fig. 2.  Basis magnitude structure: |F| (DFT, left), |U| (canonical RFT, "
        "center), and difference (right) for N = 32. The DFT has constant magnitude "
        "1/\u221AN; the RFT exhibits non-uniform structure confirming Theorem 6."),
    3: ("fig3_hurwitz_convergence_600.png",
        "Fig. 3.  Hurwitz irrationality bound verification. Fibonacci convergents "
        "p/q of \u03C6 show q\u00B2|\u03C6 \u2212 p/q| converging to 1/\u221A5 "
        "(orange dashed), the optimal constant."),
    4: ("fig4_three_distance_600.png",
        "Fig. 4.  Three-distance theorem (Lemma 4): N = 89 golden-ratio points on "
        "[0, 1) exhibit exactly 3 distinct gap lengths, confirming Steinhaus\u2013S\u00F3s."),
}

# ============================================================================
# Unicode symbol replacements for LaTeX
# ============================================================================
LATEX_SYMBOLS = {
    r'\varphi': '\u03C6', r'\phi': '\u03C6', r'\Phi': '\u03A6',
    r'\Psi': '\u03A8', r'\psi': '\u03C8',
    r'\varepsilon': '\u03B5', r'\epsilon': '\u03B5',
    r'\tau': '\u03C4', r'\sigma': '\u03C3', r'\Sigma': '\u03A3',
    r'\pi': '\u03C0', r'\Pi': '\u03A0',
    r'\alpha': '\u03B1', r'\beta': '\u03B2', r'\gamma': '\u03B3',
    r'\delta': '\u03B4', r'\Delta': '\u0394',
    r'\lambda': '\u03BB', r'\Lambda': '\u039B',
    r'\mu': '\u03BC', r'\nu': '\u03BD', r'\kappa': '\u03BA',
    r'\theta': '\u03B8', r'\rho': '\u03C1', r'\eta': '\u03B7',
    r'\omega': '\u03C9', r'\Omega': '\u03A9',
    r'\infty': '\u221E', r'\cdot': '\u00B7', r'\cdots': '\u22EF',
    r'\ldots': '\u2026', r'\times': '\u00D7', r'\leq': '\u2264',
    r'\geq': '\u2265', r'\neq': '\u2260', r'\approx': '\u2248',
    r'\equiv': '\u2261', r'\propto': '\u221D', r'\sim': '~',
    r'\in': '\u2208', r'\subset': '\u2282', r'\sum': '\u2211',
    r'\prod': '\u220F', r'\int': '\u222B',
    r'\rightarrow': '\u2192', r'\leftarrow': '\u2190',
    r'\Rightarrow': '\u21D2', r'\forall': '\u2200', r'\exists': '\u2203',
    r'\partial': '\u2202', r'\nabla': '\u2207',
    r'\lfloor': '\u230A', r'\rfloor': '\u230B',
    r'\lceil': '\u2308', r'\rceil': '\u2309',
    r'\odot': '\u2299', r'\oplus': '\u2295',
    r'\star': '\u22C6', r'\circ': '\u2218',
    r'\dagger': '\u2020', r'\ddagger': '\u2021',
    r'\ell': '\u2113', r'\hbar': '\u210F',
    r'\mathbb{R}': '\u211D', r'\mathbb{C}': '\u2102',
    r'\mathbb{Z}': '\u2124', r'\mathbb{N}': '\u2115',
    r'\square': '\u25A1', r'\qed': '\u25A1',
}


def clean_latex(text):
    """Convert LaTeX notation to readable Unicode."""
    t = text
    # Strip $$ and $
    t = re.sub(r'\$\$(.+?)\$\$', r'\1', t, flags=re.DOTALL)
    t = re.sub(r'\$(.+?)\$', r'\1', t)
    # Bold/italic markdown
    t = re.sub(r'\*{2,3}(.+?)\*{2,3}', r'\1', t)
    t = re.sub(r'(?<!\*)\*(?!\*)(.+?)(?<!\*)\*(?!\*)', r'\1', t)
    # Links
    t = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', t)
    # Code backticks
    t = re.sub(r'`([^`]+)`', r'\1', t)
    # LaTeX commands
    t = re.sub(r'\\text\{([^}]*)\}', r'\1', t)
    t = re.sub(r'\\mathrm\{([^}]*)\}', r'\1', t)
    t = re.sub(r'\\mathcal\{([^}]*)\}', r'\1', t)
    t = re.sub(r'\\mathbf\{([^}]*)\}', r'\1', t)
    t = re.sub(r'\\textstyle', '', t)
    t = re.sub(r'\\displaystyle', '', t)
    t = re.sub(r'\\frac\{([^}]*)\}\{([^}]*)\}', r'(\1)/(\2)', t)
    # sqrt with Unicode
    t = re.sub(r'\\sqrt\{([^}]*)\}', lambda m: '\u221A(' + m.group(1) + ')', t)
    # \left \right etc
    for cmd in [r'\left', r'\right', r'\bigl', r'\bigr', r'\Bigl', r'\Bigr',
                r'\big', r'\Big', r'\quad', r'\qquad', r'\,', r'\;', r'\!']:
        t = t.replace(cmd, ' ' if 'quad' in cmd else '')
    # Greek/symbols
    for cmd, uni in sorted(LATEX_SYMBOLS.items(), key=lambda x: -len(x[0])):
        t = t.replace(cmd, uni)
    # Remaining backslash commands
    t = re.sub(r'\\([a-zA-Z]+)', r'\1', t)
    # Clean up multiple spaces
    t = re.sub(r'  +', ' ', t)
    return t.strip()


# ============================================================================
# OMML equation builder
# ============================================================================
OMML_NS = "http://schemas.openxmlformats.org/officeDocument/2006/math"
WNS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

def _m(tag): return f"{{{OMML_NS}}}{tag}"
def _w(tag): return f"{{{WNS}}}{tag}"

def make_omml_display(text):
    """Create a display-mode OMML equation paragraph element."""
    oMathPara = etree.Element(_m("oMathPara"))
    oMath = etree.SubElement(oMathPara, _m("oMath"))
    r = etree.SubElement(oMath, _m("r"))
    # math run properties
    rpr = etree.SubElement(r, _m("rPr"))
    sty = etree.SubElement(rpr, _m("sty"))
    sty.set(_m("val"), "p")
    # word run properties
    wrpr = etree.SubElement(r, _w("rPr"))
    rFonts = etree.SubElement(wrpr, _w("rFonts"))
    rFonts.set(_w("ascii"), "Cambria Math")
    rFonts.set(_w("hAnsi"), "Cambria Math")
    sz = etree.SubElement(wrpr, _w("sz"))
    sz.set(_w("val"), "18")  # 9pt in half-points
    # text
    t_el = etree.SubElement(r, _m("t"))
    t_el.text = clean_latex(text)
    t_el.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    return oMathPara


# ============================================================================
# Document helpers
# ============================================================================

def clear_body(doc):
    """Remove all paragraphs and tables, keep section properties."""
    body = doc.element.body
    for child in list(body):
        tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
        if tag in ("p", "tbl"):
            body.remove(child)


def set_run_font(run, size=BODY_FONT_SIZE, name=BODY_FONT, bold=False, italic=False):
    """Configure a run's font properties."""
    run.font.size = Pt(size)
    run.font.name = name
    run.bold = bold
    run.italic = italic
    # Force font for non-Latin scripts too
    rpr = run._element.get_or_add_rPr()
    rFonts = rpr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rpr.insert(0, rFonts)
    rFonts.set(qn('w:ascii'), name)
    rFonts.set(qn('w:hAnsi'), name)
    rFonts.set(qn('w:cs'), name)


def add_body_para(doc, text, style="Normal", bold=False, italic=False,
                  font_size=BODY_FONT_SIZE, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                  space_before=0, space_after=3, first_indent=0.15):
    """Add a properly formatted body paragraph."""
    p = doc.add_paragraph(style=style)
    if text:
        run = p.add_run(text)
        set_run_font(run, size=font_size, bold=bold, italic=italic)
    p.alignment = align
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    if first_indent and first_indent > 0:
        pf.first_line_indent = Inches(first_indent)
    else:
        pf.first_line_indent = Inches(0)
    # Keep lines together
    pf.widow_control = True
    return p


def add_heading1(doc, text, style="Heading 1"):
    """Add an IEEE-style section heading (roman numeral)."""
    p = doc.add_paragraph(style=style)
    run = p.add_run(clean_latex(text))
    pf = p.paragraph_format
    pf.space_before = Pt(10)
    pf.space_after = Pt(4)
    pf.first_line_indent = Inches(0)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return p


def add_heading2(doc, text, style="Heading 2"):
    """Add an IEEE-style subsection heading."""
    p = doc.add_paragraph(style=style)
    run = p.add_run(clean_latex(text))
    pf = p.paragraph_format
    pf.space_before = Pt(6)
    pf.space_after = Pt(2)
    pf.first_line_indent = Inches(0)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return p


def add_equation(doc, latex_text):
    """Add a centered display equation as an OMML math object."""
    p = doc.add_paragraph(style="Normal")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.space_before = Pt(4)
    pf.space_after = Pt(4)
    pf.first_line_indent = Inches(0)
    omml = make_omml_display(latex_text)
    p._element.append(omml)
    return p


def add_figure(doc, fig_num):
    """Add a figure sized to fit within a single column (3.4 in)."""
    fname, caption = FIGS[fig_num]
    img_path = FIG_DIR / fname
    
    # Figure paragraph
    p = doc.add_paragraph(style="Normal")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.space_before = Pt(6)
    pf.space_after = Pt(2)
    pf.first_line_indent = Inches(0)
    
    run = p.add_run()
    if img_path.exists():
        # All figures fit within column width (3.4 in)
        run.add_picture(str(img_path), width=Inches(COL_WIDTH_IN))
    else:
        run.add_text(f"[Figure: {fname}]")
        set_run_font(run, size=9, italic=True)
    
    # Caption paragraph
    cap_p = doc.add_paragraph(style="Normal")
    cap_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    cap_pf = cap_p.paragraph_format
    cap_pf.space_before = Pt(2)
    cap_pf.space_after = Pt(6)
    cap_pf.first_line_indent = Inches(0)
    
    # Split "Fig. N." from rest
    dot_pos = caption.find('.', caption.find('.') + 1)  # second dot
    if dot_pos > 0:
        label = caption[:dot_pos + 1]
        rest = caption[dot_pos + 1:]
    else:
        label = caption
        rest = ""
    
    run_label = cap_p.add_run(label)
    set_run_font(run_label, size=CAPTION_FONT_SIZE, bold=True)
    if rest:
        run_rest = cap_p.add_run(rest)
        set_run_font(run_rest, size=CAPTION_FONT_SIZE)
    
    return p


def add_table(doc, header, rows, caption=None):
    """Add a table constrained to column width with IEEE formatting."""
    # Caption above (IEEE style)
    if caption:
        cap_p = doc.add_paragraph(style="Normal")
        cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap_pf = cap_p.paragraph_format
        cap_pf.space_before = Pt(6)
        cap_pf.space_after = Pt(2)
        cap_pf.first_line_indent = Inches(0)
        
        # Split label from description
        m = re.match(r'(TABLE\s+[IVX]+\.?)', caption)
        if m:
            run_l = cap_p.add_run(m.group(1) + " ")
            set_run_font(run_l, size=CAPTION_FONT_SIZE, bold=True)
            run_r = cap_p.add_run(caption[m.end():].strip())
            set_run_font(run_r, size=CAPTION_FONT_SIZE)
        else:
            run = cap_p.add_run(caption)
            set_run_font(run, size=CAPTION_FONT_SIZE)
    
    ncols = len(header)
    table = doc.add_table(rows=1, cols=ncols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Set table width to fit within column
    tbl = table._tbl
    tbl_pr = tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")
    
    # Table width in twips (3.4 in * 1440 twips/in)
    tbl_w = OxmlElement("w:tblW")
    tbl_w.set(qn("w:w"), str(int(COL_WIDTH_IN * 1440)))
    tbl_w.set(qn("w:type"), "dxa")
    # Remove existing tblW if any
    existing_w = tbl_pr.find(qn("w:tblW"))
    if existing_w is not None:
        tbl_pr.remove(existing_w)
    tbl_pr.insert(0, tbl_w)
    
    # Column widths
    col_w_twips = int(COL_WIDTH_IN * 1440 / ncols)
    tbl_grid = tbl.find(qn("w:tblGrid"))
    if tbl_grid is None:
        tbl_grid = OxmlElement("w:tblGrid")
        tbl.insert(1, tbl_grid)
    # Clear existing
    for gc in list(tbl_grid):
        tbl_grid.remove(gc)
    for _ in range(ncols):
        gc = OxmlElement("w:gridCol")
        gc.set(qn("w:w"), str(col_w_twips))
        tbl_grid.append(gc)
    
    # Header row
    for i, txt in enumerate(header):
        cell = table.rows[0].cells[i]
        cell.text = ""
        cp = cell.paragraphs[0]
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cp.add_run(clean_latex(txt))
        set_run_font(run, size=TABLE_FONT_SIZE, bold=True)
        # Cell margins
        _set_cell_margins(cell, top=30, bottom=30, left=30, right=30)
    
    # Data rows
    for row_data in rows:
        row = table.add_row()
        for i, txt in enumerate(row_data):
            if i < ncols:
                cell = row.cells[i]
                cell.text = ""
                cp = cell.paragraphs[0]
                cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = cp.add_run(clean_latex(txt))
                set_run_font(run, size=TABLE_FONT_SIZE)
                _set_cell_margins(cell, top=20, bottom=20, left=30, right=30)
    
    # IEEE table borders: top, bottom, header-bottom only (clean look)
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "bottom"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "6")  # 0.75pt
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "000000")
        borders.append(el)
    # Thin inside horizontal
    ih = OxmlElement("w:insideH")
    ih.set(qn("w:val"), "single")
    ih.set(qn("w:sz"), "2")  # 0.25pt  
    ih.set(qn("w:space"), "0")
    ih.set(qn("w:color"), "808080")
    borders.append(ih)
    
    existing_borders = tbl_pr.find(qn("w:tblBorders"))
    if existing_borders is not None:
        tbl_pr.remove(existing_borders)
    tbl_pr.append(borders)
    
    return table


def _set_cell_margins(cell, top=0, bottom=0, left=0, right=0):
    """Set cell margins in twips."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    mar = OxmlElement("w:tcMar")
    for side, val in [("top", top), ("bottom", bottom), ("start", left), ("end", right)]:
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:w"), str(val))
        el.set(qn("w:type"), "dxa")
        mar.append(el)
    tcPr.append(mar)


def add_reference(doc, text, style="References"):
    """Add a reference entry."""
    p = doc.add_paragraph(style=style)
    run = p.add_run(clean_latex(text))
    set_run_font(run, size=REF_FONT_SIZE)
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(1)
    pf.first_line_indent = Inches(0)
    # Hanging indent for references
    pf.left_indent = Inches(0.25)
    pf.first_line_indent = Inches(-0.25)
    return p


# ============================================================================
# Parse the markdown paper
# ============================================================================

def parse_paper(path):
    """Parse the paper into structured blocks."""
    lines = path.read_text(encoding="utf-8").splitlines()
    blocks = []
    i = 0
    in_refs = False

    while i < len(lines):
        line = lines[i]

        # Code fences
        if line.strip().startswith("```"):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            i += 1
            if code_lines:
                blocks.append({"type": "code", "text": "\n".join(code_lines)})
            continue

        # Blank / HR
        if not line.strip() or re.match(r'^-{3,}\s*$', line.strip()):
            i += 1
            continue

        # Title
        if line.startswith("# ") and not blocks:
            blocks.append({"type": "title", "text": line[2:].strip()})
            i += 1
            continue

        # Author
        if line.startswith("**") and len(blocks) <= 2:
            blocks.append({"type": "author", "text": line.strip("* \n")})
            i += 1
            continue

        # Display equation $$...$$
        if line.strip().startswith("$$"):
            eq_lines = []
            if line.strip() == "$$":
                i += 1
                while i < len(lines) and not lines[i].strip().endswith("$$"):
                    eq_lines.append(lines[i].strip())
                    i += 1
                if i < len(lines):
                    last = lines[i].strip()
                    if last != "$$":
                        eq_lines.append(last.rstrip("$"))
                i += 1
            else:
                eq_lines = [line.strip().strip("$")]
                if not line.strip().endswith("$$") or line.strip().count("$$") < 2:
                    i += 1
                    while i < len(lines) and not lines[i].strip().endswith("$$"):
                        eq_lines.append(lines[i].strip())
                        i += 1
                    if i < len(lines):
                        eq_lines.append(lines[i].strip().rstrip("$"))
                i += 1
            blocks.append({"type": "equation", "text": " ".join(eq_lines)})
            continue

        # H1 (##)
        m1 = re.match(r'^##\s+(.+)', line)
        if m1:
            htxt = m1.group(1).strip()
            if htxt.lower().startswith("references"):
                in_refs = True
            blocks.append({"type": "h1", "text": htxt})
            i += 1
            continue

        # H2 (###)
        m2 = re.match(r'^###\s+(.+)', line)
        if m2:
            blocks.append({"type": "h2", "text": m2.group(1).strip()})
            i += 1
            continue

        # Abstract
        if line.strip().startswith("*Abstract*") or line.strip().startswith("Abstract"):
            blocks.append({"type": "abstract", "text": line.strip()})
            i += 1
            continue

        # Index terms
        if line.strip().startswith("*Index Terms*") or line.strip().startswith("Index Terms"):
            blocks.append({"type": "index_terms", "text": line.strip()})
            i += 1
            continue

        # Table
        if "|" in line and i + 1 < len(lines) and re.match(r'^\|[\s\-:|]+\|', lines[i + 1]):
            hdr = [c.strip() for c in line.strip().strip("|").split("|")]
            i += 2
            rows = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                row = [c.strip() for c in lines[i].strip().strip("|").split("|")]
                rows.append(row)
                i += 1
            blocks.append({"type": "table", "header": hdr, "rows": rows})
            continue

        # Reference entry
        if in_refs and re.match(r'^\[(\d+)\]', line.strip()):
            blocks.append({"type": "ref", "text": line.strip()})
            i += 1
            continue

        # Normal paragraph
        para_lines = [line]
        i += 1
        while i < len(lines):
            nl = lines[i]
            if not nl.strip() or nl.startswith("#") or nl.strip().startswith("$$"):
                break
            if re.match(r'^-{3,}', nl.strip()):
                break
            if "|" in nl and i + 1 < len(lines) and re.match(r'^\|[\s\-:|]+\|', lines[i + 1]):
                break
            if in_refs and re.match(r'^\[(\d+)\]', nl.strip()):
                break
            para_lines.append(nl)
            i += 1
        full = " ".join(l.strip() for l in para_lines)
        blocks.append({"type": "normal", "text": full})

    return blocks


# ============================================================================
# Figure placement logic
# ============================================================================

# Map: figure number → keyword to look for in text (insert figure after that paragraph)
FIG_TRIGGERS = {
    1: "Fig. 1",       # mentioned in worked example
    3: "Hurwitz",      # after Hurwitz lemma
    4: "three-distance",  # after three-distance lemma
}

TABLE_CAPTIONS = [
    "TABLE I. Spectral concentration comparison",
    "TABLE II. Coefficient decay comparison (N = 256, mean over 500 phase draws, seed 42)",
    "TABLE III. FPGA synthesis results (16\u00D716, Q1.15)",
    "TABLE IV. Per-harmonic DFT leakage (Lemma 7)",
    "TABLE V. Per-harmonic comparison (Lemma 8)",
    "TABLE VI. Ensemble concentration gap (Lemma 9), 200 trials per size",
    "TABLE VII. Entropic uncertainty verification (N = 64)",
]


# ============================================================================
# Main build
# ============================================================================

def main():
    print(f"Loading template: {TEMPLATE.name}")
    doc = Document(str(TEMPLATE))

    style_names = {s.name for s in doc.styles}
    def ss(preferred, fallback="Normal"):
        return preferred if preferred in style_names else fallback

    print("Clearing body...")
    clear_body(doc)

    print("Parsing paper...")
    blocks = parse_paper(MD_FILE)
    print(f"  {len(blocks)} blocks")

    table_idx = 0
    figs_inserted = set()

    for bi, blk in enumerate(blocks):
        btype = blk["type"]
        raw = blk.get("text", "")
        txt = clean_latex(raw) if raw else ""

        # ── Title ──
        if btype == "title":
            p = doc.add_paragraph(style=ss("Title"))
            run = p.add_run(txt)
            set_run_font(run, size=24, name="Times New Roman")
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pf = p.paragraph_format
            pf.space_before = Pt(12)
            pf.space_after = Pt(6)
            pf.first_line_indent = Inches(0)

        # ── Author ──
        elif btype == "author":
            p = doc.add_paragraph(style=ss("Normal"))
            run = p.add_run("Luis Michael Minier")
            set_run_font(run, size=11, name="Times New Roman")
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(10)
            p.paragraph_format.first_line_indent = Inches(0)

        # ── Abstract ──
        elif btype == "abstract":
            p = doc.add_paragraph(style=ss("Normal"))
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
            p.paragraph_format.first_line_indent = Inches(0)
            
            body = re.sub(r'^\*?Abstract\*?\s*', '', raw).lstrip('\u2014 -')
            run_b = p.add_run("Abstract\u2014")
            set_run_font(run_b, size=9, bold=True)
            run_t = p.add_run(clean_latex(body))
            set_run_font(run_t, size=9)

        # ── Index Terms ──
        elif btype == "index_terms":
            p = doc.add_paragraph(style=ss("Normal"))
            p.paragraph_format.space_after = Pt(8)
            p.paragraph_format.first_line_indent = Inches(0)
            
            body = re.sub(r'^\*?Index Terms\*?\s*', '', raw).lstrip('\u2014 -')
            run_b = p.add_run("Index Terms\u2014")
            set_run_font(run_b, size=9, bold=True, italic=True)
            run_t = p.add_run(clean_latex(body))
            set_run_font(run_t, size=9, italic=True)

        # ── Heading 1 ──
        elif btype == "h1":
            add_heading1(doc, raw, ss("Heading 1"))

        # ── Heading 2 ──
        elif btype == "h2":
            add_heading2(doc, raw, ss("Heading 2"))

        # ── Display equation ──
        elif btype == "equation":
            add_equation(doc, raw)

        # ── Table ──
        elif btype == "table":
            cap = TABLE_CAPTIONS[table_idx] if table_idx < len(TABLE_CAPTIONS) else None
            add_table(doc, blk["header"], blk["rows"], caption=cap)
            table_idx += 1

        # ── Reference ──
        elif btype == "ref":
            add_reference(doc, raw, ss("References", "Normal"))

        # ── Code ──
        elif btype == "code":
            p = add_body_para(doc, txt, font_size=7, align=WD_ALIGN_PARAGRAPH.LEFT,
                              first_indent=0)
            for run in p.runs:
                run.font.name = "Courier New"

        # ── Normal paragraph ──
        else:
            # Skip standalone table caption lines (handled with table)
            if re.match(r'^\*?\*?TABLE\s+[IVX]+', raw):
                continue

            # Detect theorem/definition/lemma labels
            label_re = re.match(
                r'\*?\*?(Theorem \d+|Definition \d+|Lemma \d+|Corollary \d+\.?\d*|'
                r'Proposition \d+|Proof\.?|Proof \(|Proof:)\s*',
                txt
            )
            
            if label_re:
                p = doc.add_paragraph(style=ss("Normal"))
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                pf = p.paragraph_format
                pf.space_before = Pt(3)
                pf.space_after = Pt(3)
                pf.first_line_indent = Inches(0)
                
                label = label_re.group(1).rstrip(".")
                rest = txt[label_re.end():].lstrip(". :)")
                
                run_l = p.add_run(label + ". ")
                set_run_font(run_l, size=BODY_FONT_SIZE, bold=True, italic=True)
                
                if rest:
                    run_r = p.add_run(rest)
                    set_run_font(run_r, size=BODY_FONT_SIZE)
            else:
                add_body_para(doc, txt, first_indent=0.15, space_after=3)

        # ── Insert figures at appropriate positions ──
        for fig_num, trigger in FIG_TRIGGERS.items():
            if fig_num not in figs_inserted and trigger.lower() in raw.lower():
                add_figure(doc, fig_num)
                figs_inserted.add(fig_num)

    # Insert Fig 2 after section III (foundational properties)
    # and any remaining figures at the end
    remaining_figs = sorted(set(FIGS.keys()) - figs_inserted)
    for fig_num in remaining_figs:
        add_figure(doc, fig_num)

    print(f"Saving: {TEMPLATE.name}")
    doc.save(str(TEMPLATE))
    
    # Summary
    print(f"\n  Paragraphs: {len(doc.paragraphs)}")
    print(f"  Tables:     {len(doc.tables)}")
    img_count = sum(1 for rel in doc.part.rels.values() if "image" in rel.reltype)
    print(f"  Images:     {img_count}")
    
    eq_ns = "http://schemas.openxmlformats.org/officeDocument/2006/math"
    eq_count = len(doc.element.body.findall(f'.//{{{eq_ns}}}oMathPara'))
    print(f"  Equations:  {eq_count}")
    print("Done!")


if __name__ == "__main__":
    main()
