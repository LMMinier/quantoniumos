#!/usr/bin/env python3
"""
Populate the IEEE Transactions Word template with the RFT Diophantine paper.

Reads:  papers/RFT_DIOPHANTINE_PAPER.md
Uses:   papers/Transactions-template-and-instructions-on-how-to-create-your-article-formatted (4).docx  (as style donor)
Writes: papers/RFT_Diophantine_IEEE.docx
"""

import copy, re, textwrap
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path(__file__).resolve().parent.parent
TEMPLATE = ROOT / "papers" / "Transactions-template-and-instructions-on-how-to-create-your-article-formatted (4).docx"
MD_FILE  = ROOT / "papers" / "RFT_DIOPHANTINE_PAPER.md"
OUTPUT   = ROOT / "papers" / "RFT_Diophantine_IEEE.docx"

# ---------------------------------------------------------------------------
# helpers
# ---------------------------------------------------------------------------

def clear_body(doc):
    """Remove every paragraph and table from the document body."""
    body = doc.element.body
    for child in list(body):
        tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
        if tag in ("p", "tbl", "sectPr"):
            if tag == "sectPr":
                continue          # keep section properties (columns, margins)
            body.remove(child)


def add_paragraph(doc, text, style_name, bold=False, italic=False, alignment=None, font_size=None):
    """Append a paragraph with the given style.  Returns the paragraph."""
    p = doc.add_paragraph(style=style_name)
    run = p.add_run(text)
    if bold:
        run.bold = True
    if italic:
        run.italic = True
    if alignment is not None:
        p.alignment = alignment
    if font_size is not None:
        run.font.size = Pt(font_size)
    return p


def add_table(doc, header_row, data_rows):
    """Add a simple table matching IEEE style (no fancy shading)."""
    ncols = len(header_row)
    table = doc.add_table(rows=1, cols=ncols)
    table.style = "Normal Table"
    # header
    for i, txt in enumerate(header_row):
        cell = table.rows[0].cells[i]
        cell.text = txt
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(8)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    # data
    for row_data in data_rows:
        row = table.add_row()
        for i, txt in enumerate(row_data):
            if i < ncols:
                cell = row.cells[i]
                cell.text = txt
                for run in cell.paragraphs[0].runs:
                    run.font.size = Pt(8)
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    # thin borders
    tbl = table._tbl
    tbl_pr = tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "000000")
        borders.append(el)
    tbl_pr.append(borders)
    return table


def strip_dollars(txt):
    """Remove $ wrappers for plain-text Word (no LaTeX rendering)."""
    # keep the math content, just strip dollar signs
    txt = re.sub(r'\$\$(.+?)\$\$', r'\1', txt, flags=re.DOTALL)
    txt = re.sub(r'\$(.+?)\$', r'\1', txt)
    return txt


# ---------------------------------------------------------------------------
# Markdown → structured blocks
# ---------------------------------------------------------------------------

def parse_md(path):
    """
    Very lightweight Markdown parser.  Returns a list of dicts:
      {type: title|h1|h2|abstract|index_terms|normal|table_caption|table|ref, text:…}
    """
    lines = path.read_text(encoding="utf-8").splitlines()
    blocks = []
    i = 0
    in_references = False
    in_code_block = False

    while i < len(lines):
        line = lines[i]

        # skip code fences
        if line.strip().startswith("```"):
            in_code_block = not in_code_block
            if not in_code_block:
                i += 1
                continue
            # start of code block – collect until close
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            in_code_block = False
            blocks.append({"type": "code", "text": "\n".join(code_lines)})
            i += 1
            continue

        # blank
        if not line.strip():
            i += 1
            continue

        # horizontal rule
        if re.match(r'^-{3,}\s*$', line.strip()):
            i += 1
            continue

        # title (# at top)
        if line.startswith("# ") and len(blocks) == 0:
            blocks.append({"type": "title", "text": line[2:].strip()})
            i += 1
            continue

        # heading 1: ## with roman numeral or "References" or "Appendix"
        m1 = re.match(r'^##\s+(.+)', line)
        if m1:
            htxt = m1.group(1).strip()
            if htxt.lower().startswith("references"):
                in_references = True
            blocks.append({"type": "h1", "text": htxt})
            i += 1
            continue

        # heading 2: ###
        m2 = re.match(r'^###\s+(.+)', line)
        if m2:
            blocks.append({"type": "h2", "text": m2.group(1).strip()})
            i += 1
            continue

        # bold author line (right after title)
        if line.startswith("**") and len(blocks) <= 2:
            blocks.append({"type": "author", "text": line.strip().strip("*")})
            i += 1
            continue

        # abstract line
        if line.strip().startswith("*Abstract*") or line.strip().startswith("Abstract"):
            blocks.append({"type": "abstract", "text": line.strip().lstrip("*").lstrip()})
            i += 1
            continue

        # index terms
        if line.strip().startswith("*Index Terms*") or line.strip().startswith("Index Terms"):
            blocks.append({"type": "index_terms", "text": line.strip().lstrip("*").lstrip()})
            i += 1
            continue

        # table: detect | header | ... | followed by |---|
        if "|" in line and i + 1 < len(lines) and re.match(r'^\|[\s\-:|]+\|', lines[i + 1]):
            # parse header
            hdr = [c.strip() for c in line.strip().strip("|").split("|")]
            i += 2  # skip separator
            rows = []
            while i < len(lines) and "|" in lines[i] and lines[i].strip().startswith("|"):
                row = [c.strip() for c in lines[i].strip().strip("|").split("|")]
                rows.append(row)
                i += 1
            blocks.append({"type": "table", "header": hdr, "rows": rows})
            continue

        # reference entries [N]
        if in_references and re.match(r'^\[(\d+)\]', line.strip()):
            blocks.append({"type": "ref", "text": line.strip()})
            i += 1
            continue

        # regular paragraph (may be multi-line until blank)
        para_lines = [line]
        i += 1
        while i < len(lines) and lines[i].strip() and not lines[i].startswith("#") and not lines[i].startswith("---") and not ("|" in lines[i] and i + 1 < len(lines) and re.match(r'^\|[\s\-:|]+\|', lines[i + 1])):
            if re.match(r'^\[(\d+)\]', lines[i].strip()) and in_references:
                break
            para_lines.append(lines[i])
            i += 1
        full = " ".join(l.strip() for l in para_lines)
        blocks.append({"type": "normal", "text": full})

    return blocks


# ---------------------------------------------------------------------------
# main
# ---------------------------------------------------------------------------

def main():
    print(f"Loading template: {TEMPLATE.name}")
    doc = Document(str(TEMPLATE))

    print("Clearing template sample content …")
    clear_body(doc)

    print(f"Parsing paper: {MD_FILE.name}")
    blocks = parse_md(MD_FILE)
    print(f"  → {len(blocks)} blocks")

    # Verify styles exist
    style_names = {s.name for s in doc.styles}
    def safe_style(preferred, fallback="Normal"):
        return preferred if preferred in style_names else fallback

    for blk in blocks:
        txt = strip_dollars(blk["text"]) if "text" in blk else ""
        # strip residual markdown bold/italic markers for Word
        txt = re.sub(r'\*{2,3}(.+?)\*{2,3}', r'\1', txt)
        txt = re.sub(r'(?<!\*)\*(?!\*)(.+?)(?<!\*)\*(?!\*)', r'\1', txt)
        # strip markdown link syntax
        txt = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', txt)

        btype = blk["type"]

        if btype == "title":
            add_paragraph(doc, txt, safe_style("Title"), alignment=WD_ALIGN_PARAGRAPH.CENTER)

        elif btype == "author":
            add_paragraph(doc, txt, safe_style("Normal"),
                          alignment=WD_ALIGN_PARAGRAPH.CENTER, font_size=11)

        elif btype == "abstract":
            # IEEE abstract: bold "Abstract—" then normal text
            p = doc.add_paragraph(style=safe_style("Normal"))
            if txt.startswith("Abstract"):
                dash_pos = txt.find("—")
                if dash_pos == -1:
                    dash_pos = txt.find("-")
                if dash_pos > 0:
                    run_b = p.add_run(txt[:dash_pos + 1])
                    run_b.bold = True
                    run_b.font.size = Pt(9)
                    run_n = p.add_run(txt[dash_pos + 1:])
                    run_n.font.size = Pt(9)
                else:
                    run = p.add_run(txt)
                    run.font.size = Pt(9)
            else:
                p.add_run(txt).font.size = Pt(9)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        elif btype == "index_terms":
            p = doc.add_paragraph(style=safe_style("Normal"))
            if "—" in txt:
                dash = txt.index("—")
                run_b = p.add_run(txt[:dash + 1])
                run_b.bold = True
                run_b.italic = True
                run_b.font.size = Pt(9)
                run_n = p.add_run(txt[dash + 1:])
                run_n.font.size = Pt(9)
            else:
                p.add_run(txt).font.size = Pt(9)

        elif btype == "h1":
            add_paragraph(doc, txt.upper() if not txt[0].isdigit() and "Appendix" not in txt else txt,
                          safe_style("Heading 1"))

        elif btype == "h2":
            add_paragraph(doc, txt, safe_style("Heading 2"))

        elif btype == "table":
            add_table(doc, blk["header"], blk["rows"])

        elif btype == "ref":
            add_paragraph(doc, txt, safe_style("References", "Normal"), font_size=8)

        elif btype == "code":
            # Code block → small monospace paragraph
            p = doc.add_paragraph(style=safe_style("Normal"))
            run = p.add_run(txt)
            run.font.size = Pt(7)
            run.font.name = "Courier New"

        else:  # normal
            add_paragraph(doc, txt, safe_style("Normal"), font_size=10)

    doc.save(str(OUTPUT))
    print(f"\nSaved → {OUTPUT}")
    print(f"  paragraphs: {len(doc.paragraphs)}")


if __name__ == "__main__":
    main()
